"""简历相关路由"""
import os
from typing import Optional
from fastapi import APIRouter, UploadFile, File, Form, Depends
from pydantic import BaseModel
from auth_deps import get_current_session_id

router = APIRouter()

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'data', 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)


class ResumeEdit(BaseModel):
    name: Optional[str] = ""
    education_background: Optional[list] = []
    skills: Optional[list] = []
    work_years: Optional[int] = None
    certificates: Optional[list] = []
    work_experience: Optional[list] = []
    projects: Optional[list] = []
    internships: Optional[list] = []


def _extract_text(file_bytes: bytes, filename: str) -> str:
    """从PDF/DOCX提取文本（复用前原型逻辑）"""
    ext = filename.rsplit('.', 1)[-1].lower()
    from io import BytesIO

    if ext == 'pdf':
        try:
            import pdfplumber
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                texts = [p.extract_text() for p in pdf.pages if p.extract_text()]
            return '\n'.join(texts) if texts else ""
        except Exception:
            pass
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(BytesIO(file_bytes))
            return '\n'.join(p.extract_text() or '' for p in reader.pages)
        except Exception:
            return ""

    if ext == 'docx':
        try:
            from docx import Document
            import zipfile, re
            doc = Document(BytesIO(file_bytes))
            texts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text.strip())
            if not texts:
                with zipfile.ZipFile(BytesIO(file_bytes)) as z:
                    xml = z.read('word/document.xml').decode('utf-8')
                found = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml)
                if found:
                    texts.append(''.join(found))
            return '\n'.join(texts)
        except Exception:
            return ""

    return ""


@router.post("/upload")
async def upload_resume(file: UploadFile = File(...), session_id: str = Depends(get_current_session_id)):
    """上传简历并解析"""
    from database import save_resume
    from llm_service import parse_resume

    try:
        content = await file.read()
        text = _extract_text(content, file.filename)

        if not text:
            return {"ok": False, "error": "无法读取文件内容，请确认上传的是文本版PDF/Word"}

        parsed = parse_resume(text)
        save_resume(session_id, parsed)
        return {"ok": True, "data": parsed}
    except Exception as e:
        print(f"简历上传失败: {e}")
        return {"ok": False, "error": f"简历解析失败，请重试: {str(e)}"}


@router.get("/get")
def get_resume(session_id: str = Depends(get_current_session_id)):
    """获取已保存的简历"""
    from database import get_resume as db_get
    try:
        data = db_get(session_id)
        return {"ok": True, "data": data}
    except Exception as e:
        print(f"获取简历失败: {e}")
        return {"ok": False, "error": str(e)}


@router.post("/update")
def update_resume(data: ResumeEdit, session_id: str = Depends(get_current_session_id)):
    """手动更新简历字段"""
    from database import save_resume
    from database import get_resume as db_get
    try:
        existing = db_get(session_id) or {}
        merged = {
            **existing,
            'name': data.name or existing.get('name', ''),
            'education_background': data.education_background or existing.get('education_background', []),
            'skills': data.skills or existing.get('skills', []),
            'work_years': max(0, data.work_years if data.work_years is not None else existing.get('work_years', 0)),
            'certificates': data.certificates or existing.get('certificates', []),
            'work_experience': data.work_experience or existing.get('work_experience', []),
            'projects': data.projects or existing.get('projects', []),
            'internships': data.internships or existing.get('internships', []),
            'confirmed': True,
        }
        save_resume(session_id, merged)
        return {"ok": True, "data": merged}
    except Exception as e:
        print(f"更新简历失败: {e}")
        return {"ok": False, "error": str(e)}
